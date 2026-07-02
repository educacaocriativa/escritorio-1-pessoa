"""Exporta um documento jurídico gerado para .docx (download).

Conversão markdown-leve → Word: títulos (#, ##, ###), listas (-, *, 1.) e parágrafos.
Unicode nativo (acentos PT-BR) — usa python-docx, já dependência do projeto.
"""
from __future__ import annotations

import io
import re


def to_docx(*, title: str, content: str, metadata_raw: str = "") -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(title, level=0)
    _render(doc, content)

    if metadata_raw.strip():
        doc.add_page_break()
        doc.add_heading("Metadados do Relatório", level=1)
        _render(doc, metadata_raw)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render(doc, text: str) -> None:
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif re.match(r"^\s*[-*]\s+", line):
            doc.add_paragraph(_clean(re.sub(r"^\s*[-*]\s+", "", line)), style="List Bullet")
        elif re.match(r"^\s*\d+\.\s+", line):
            doc.add_paragraph(_clean(re.sub(r"^\s*\d+\.\s+", "", line)), style="List Number")
        else:
            doc.add_paragraph(_clean(line))


def _clean(s: str) -> str:
    # remove marcadores de negrito/itálico do markdown (o Word não os interpreta como texto)
    return re.sub(r"\*\*(.+?)\*\*", r"\1", re.sub(r"(?<!\*)\*(?!\*)(.+?)\*", r"\1", s))
