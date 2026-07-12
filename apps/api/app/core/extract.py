"""Extração de texto de arquivos enviados (PDF, Word, imagem com OCR, txt).

Trabalha em bytes (em memória) — nada é gravado em disco. O texto extraído passa pelo
anonimizador antes de ir para a IA (ver anonymizer.py). Portado do lex-intelligentia-app.
"""
from __future__ import annotations

import io

# Tipos aceitos no upload do Assistente Jurídico.
ACCEPTED = (
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "image/jpeg",
    "image/png",
    "text/plain",
)


def extract_text(data: bytes, content_type: str, filename: str = "") -> str:
    """Extrai texto de um arquivo em memória. Devolve string (vazia se nada extraível)."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    try:
        if content_type == "application/pdf" or ext == "pdf":
            return _pdf(data)
        if "wordprocessingml" in content_type or content_type == "application/msword" or ext in (
            "docx",
            "doc",
        ):
            return _docx(data)
        if content_type.startswith("image/") or ext in ("jpg", "jpeg", "png", "tiff", "bmp"):
            return _image(data)
        if content_type.startswith("text/") or ext in ("txt", "md"):
            return data.decode("utf-8", errors="ignore")
    except Exception as e:  # noqa: BLE001 — extração nunca derruba a requisição
        return f"[ERRO NA EXTRAÇÃO: {e}]"
    return f"[TIPO NÃO SUPORTADO: {content_type or ext or 'desconhecido'}]"


def _pdf(data: bytes) -> str:
    import pdfplumber

    texts: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                texts.append(t)
    return "\n\n".join(texts) if texts else "[PDF sem texto extraível — tente OCR]"


def _docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            line = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if line:
                parts.append(line)
    return "\n\n".join(parts)


def _image(data: bytes) -> str:
    try:
        import pytesseract
        from PIL import Image

        text = pytesseract.image_to_string(Image.open(io.BytesIO(data)), lang="por")
        return text.strip() or "[Imagem sem texto detectável]"
    except ImportError:
        return "[OCR não disponível — tesseract não instalado]"
